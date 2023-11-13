# Configurar o driver do Selenium (neste exemplo, usaremos o Chrome)
# Modo 'Abre e Fecha' do Chrome corrigido

from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import Keys, ActionChains
from selenium.webdriver.common.by import By
from PIL import Image
from time import sleep
import pyautogui
from docx import Document
import paragraphs
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Screenshot import Screenshot

# -*- coding: utf-8 -*-



options = webdriver.ChromeOptions()
options.add_experimental_option("detach", True)
driver = webdriver.Chrome(options=options, service=Service(ChromeDriverManager().install()))

# Carregar a página da web desejada
driver.get("https://portalsas.com.br/login?redirect=https%3A%2F%2Fhome.portalsas.com.br%2F")
driver.maximize_window() # Modo Tela Cheia

# Carregar usuário/senha
pyautogui.click(x=504, y=399)
pyautogui.write("kennyo.cavalcante@colegio-evolucao.com")
pyautogui.click(x=504, y=508)
pyautogui.write("Kennyow86")

# Confirmar
driver.find_element('xpath', '//*[@id="login-form-submit"]').click()
sleep(3)

# Fechando Pop ups
WebDriverWait(driver, 30).until(
EC.presence_of_element_located(('xpath', '//*[@id="root"]/div[1]/div/div')) 
)

pyautogui.click(x=408, y=671)
pyautogui.click(x=1188, y=387)
driver.find_element('xpath', '//*[@id="root"]/nav/div[2]/button').click()

# Acessando 
driver.find_element('xpath', '//*[@id="root"]/nav/div[1]/div[1]/ul/li[6]/button').click()
driver.find_element('xpath', '//*[@id="root"]/nav/div[1]/div[1]/ul/li[6]/div/ul/li[6]/a').click()

# No Studos
sleep(12)

# Troca de Janela
window_name = driver.window_handles[0]
driver.switch_to.window(window_name=window_name)
driver.close()
sleep(2)
driver.switch_to.window(driver.window_handles[0])
sleep(5)

# Fechar o alert
xpathe = '//*[@id="main-container 3"]/div[2]/div/div[3]/div/div[2]/div/div/div/div/div[1]/button'
td_elements = driver.find_element(By.XPATH, xpathe)
actions = ActionChains(driver)
td_elements.click()
sleep(3)

# Inserindo a disciplina no box

#pyautogui.click(x=369, y=444)
'''xpath_input = '/html/body/div[2]/div[1]/div[2]/div/div[3]/div/div[2]/div/div/div/div/form/div[1]/input'
td_elements2 = driver.find_element(By.XPATH, xpath_input)
acao2 = ActionChains(driver)
acao2.move_to_element(td_elements2).perform()
sleep(3)'''


input_box = driver.find_element_by_name('input-search')

# Replace 'Your text here' with the string you want to input
text_to_input = 'MARÍÍÍA'

# Input the text into the input box
'''input_box.send_keys(text_to_input)
td_elements2.click()
pyautogui.click(td_elements2.click())'''
print("chegou4?")
sleep(1)
avali = str("A2 - QUÍMICA E BIOLOGIA - 2ª Série II Tri. 2023").upper()
'''acao2.send_keys(avali)'''
print("chegou5?")
#pyautogui.write(avali)

#driver.refresh() Reload da página que só carregava os exercícios de A2
sleep(6)


# Acessando os Resultados da atividade
pyautogui.click(x=1204, y=699)
pyautogui.click(x=1033, y=647)
sleep(7)



'''# Tirando os prints
#1
pyautogui.moveTo(x=1355, y=164)
pyautogui.mouseDown()
pyautogui.moveTo(x=1355, y=230)
driver.get_screenshot_as_file("mapa1.png")
sleep(1)
im = Image.open("mapa1.png")
im = im.crop((int(320), int(0), int(1200), int(580)))
im.save("mapa1.png")
#2
pyautogui.mouseDown()
pyautogui.moveTo(x=1355, y=300)
driver.get_screenshot_as_file("mapa2.png")
sleep(1)
im = Image.open("mapa2.png")
im = im.crop((int(320), int(0), int(1200), int(580)))
im.save("mapa2.png")
#3
pyautogui.mouseDown()
pyautogui.moveTo(x=1355, y=405)
driver.get_screenshot_as_file("mapa3.png")
sleep(1)
im = Image.open("mapa3.png")
im = im.crop((int(320), int(0), int(1200), int(580)))
im.save("mapa3.png")

document = Document()
document.add_heading(avali, 2)
document.add_paragraph("PROGRESSO ATIVIDADE")
document.add_picture("mapa1.png", width=Inches(6.50))
document.add_paragraph("ASSUNTO EM DESTAQUE")
document.add_picture("mapa2.png", width=Inches(6.50))
document.add_paragraph("ALUNOS EM DESTAQUE")
document.add_picture("mapa3.png", width=Inches(6.50))
document.save("docteste.docx")

#driver.close()'''

# Find all <td> elements with any class
sleep(3)
lista_percentagem = []
td_texts = []
'''xpathe = '/html/body/div[2]/div[1]/div[2]/div/div[3]/div/div[3]/div[2]/div[4]/div/div/div/div/div/div'
td_elements = driver.find_element(By.XPATH, xpathe)'''
xpathe = '/html/body/div[2]/div[1]/div[2]/div/div[3]/div/div[3]/div[2]/div[1]/div/div/div/ul/li[1]/a'
td_elements = driver.find_element(By.XPATH, xpathe)
actions = ActionChains(driver)
actions.move_to_element(td_elements).perform()
td_elements.click()
sleep(1)


#   Questão 1
pyautogui.click(x=397, y=564)
sleep(2)

'''xpathe2 = '//*[@id="chakra-modal--body-366"]/header/header/h3'''
td_elements = driver.find_element(By.CLASS_NAME, 'chakra-ui-light')
WebDriverWait(driver, 10).until(EC.frame_to_be_available_and_switch_to_it(td_elements))

'''iframe_locator = 'tag-manager'  # Replace 'your_iframe_id' with the actual ID of your iframe
driver.switch_to.frame(iframe_locator)'''

element_inside_iframe = driver.find_element(By.CLASS_NAME, 'chakra-ui-light')
actions = ActionChains(driver)
actions.move_to_element(element_inside_iframe).perform()
sleep(1)



# Iframe Pop Up

'''driver.find_element(By.CLASS_NAME, "ReactModal__Content ReactModal__Content--after-open").click()
sleep(1)'''



# frame = WebDriverWait(driver, 10).until(EC.frame_to_be_available_and_switch_to_it(By.ID, xpathe3))
'''actions = ActionChains(driver)
actions.move_to_element(td_elements).perform()'''

ob = Screenshot.Screenshot()
img_url = ob.full_screenshot(driver, save_path=r'.', image_name='myimage.png', is_load_at_runtime=True,
                                          load_wait_time=4)
driver.find_element(By.XPATH, '/html/body/div[5]/div/div/div/div[1]/div[2]/div/div/span/svg').click()
sleep(2)


'''for numero in range(1,31):
    #xpathe = '/html/body/div[2]/div[1]/div[2]/div/div[3]/div/div[3]/div[2]/div[2]/div/div/div/div/div/div/div[2]/div/div[1]/div[2]/div/div['+ str(numero)+']/button/svg/path[2]'
    
    num = int(td_elements.text.strip()[:-1])
    lista_percentagem.append(num)
    
    for element in td_elements:
        
    /html/body/div[2]/div[1]/div[2]/div/div[3]/div/div[3]/div[2]/div[2]/div/div/div/div/div/div/div[2]/div/div[1]/div[2]/div/div[1]/button/svg/path[2]
    /html/body/div[2]/div[1]/div[2]/div/div[3]/div/div[3]/div[2]/div[2]/div/div/div/div/div/div/div[2]/div/div[1]/div[2]/div/div[2]/button/svg/path[2]









        #print(element.text.strip())
        #print(num)
        if num < 60:
            //*[@id="questionReport"]/div/div/div/div/div/div/div/div/div[1]/table/tbody/tr[1]/td[1]
            //*[@id="questionReport"]/div/div/div/div/div/div/div/div/div[1]/table/tbody/tr[2]/td[1]
            actions = ActionChains(driver)
            actions.move_to_element(td_elements).perform()
            #amiguri.click()
          
            x_position = td_elements.location['x']
            y_position = td_elements.location['y']
            pyautogui.click(x_position, y_position)
            sleep(10)
            print(f"Width: {x_position}px")
            print(f"Height: {y_position}px")'''
'''td_elements = driver.find_elements(By.XPATH, '//*[@id="questionReport"]/div/div/div/div/div/div/div/div/div[1]/table/tbody/tr[2]/td[5]'''     


# Extract the text from each element and put them into a list
#td_texts = [element.text for element in td_elements]
#print(lista_percentagem)
# Print the list of extracted texts
#print(td_texts)

'''
# Print página inteira
ob = Screenshot.Screenshot()
img_url = ob.full_screenshot(driver, save_path=r'.', image_name='myimage.png', is_load_at_runtime=True,
                                          load_wait_time=4)'''