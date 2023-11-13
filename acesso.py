from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import Keys, ActionChains
from selenium.webdriver.common.by import By
from PIL import Image
from time import sleep
import pyautogui
from docx import Document
import paragraphs
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Configurar o driver do Selenium (neste exemplo, usaremos o Chrome)
# Modo 'Abre e Fecha' do Chrome corrigido

options = webdriver.ChromeOptions()
options.add_experimental_option("detach", True)
driver = webdriver.Chrome(options=options, service=Service(ChromeDriverManager().install()))

# Carregar a página da web desejada
driver.get("https://portalsas.com.br/login?redirect=https%3A%2F%2Fhome.portalsas.com.br%2F")
driver.maximize_window() # Modo Tela Cheia

# Carregar usuário/senha
pyautogui.click(x=504, y=399)
pyautogui.write("kennyo.cavalcante@colegio-evolucao.com")
pyautogui.click(x=504, y=508)
pyautogui.write("Kennyow86")

# Confirmar
driver.find_element('xpath', '//*[@id="login-form-submit"]').click()
sleep(3)

# Fechando Pop ups
WebDriverWait(driver, 30).until(
EC.presence_of_element_located(('xpath', '//*[@id="root"]/div[1]/div/div')) 
)

pyautogui.click(x=408, y=671)
pyautogui.click(x=1188, y=387)
driver.find_element('xpath', '//*[@id="root"]/nav/div[2]/button').click()

# Acessando 
driver.find_element('xpath', '//*[@id="root"]/nav/div[1]/div[1]/ul/li[6]/button').click()
driver.find_element('xpath', '//*[@id="root"]/nav/div[1]/div[1]/ul/li[6]/div/ul/li[6]/a').click()

# No Studos
sleep(12)

avali = str("A2 - QUÍMICA E BIOLOGIA ")

# Inserindo a disciplina no box
pyautogui.click(x=369, y=444)
pyautogui.write(avali)
sleep(8)

# Acessando os Resultados da atividade
pyautogui.click(x=1204, y=699)
pyautogui.click(x=1033, y=647)
sleep(10)

# Troca de Janela

window_name = driver.window_handles[0]
driver.switch_to.window(window_name=window_name)
driver.close()
sleep(2)
'''chwd = driver.window_handles
old = driver.current_window_handle    
driver.execute_script("window.open('{arguments[0]');")
WebDriverWait(driver,60).until(EC.number_of_windows_to_be(len(chwd)+1))
driver.switch_to.window()
driver.close()'''
driver.switch_to.window(driver.window_handles[0])
sleep(1)

# Tirando os prints
#1
pyautogui.moveTo(x=1355, y=164)
pyautogui.mouseDown()
pyautogui.moveTo(x=1355, y=230)
driver.get_screenshot_as_file("mapa1.png")
sleep(1)
im = Image.open("mapa1.png")
im = im.crop((int(320), int(0), int(1200), int(580)))
im.save("mapa1.png")
#2
pyautogui.mouseDown()
pyautogui.moveTo(x=1355, y=300)
driver.get_screenshot_as_file("mapa2.png")
sleep(1)
im = Image.open("mapa2.png")
im = im.crop((int(320), int(0), int(1200), int(580)))
im.save("mapa2.png")
#3
pyautogui.mouseDown()
pyautogui.moveTo(x=1355, y=405)
driver.get_screenshot_as_file("mapa3.png")
sleep(1)
im = Image.open("mapa3.png")
im = im.crop((int(320), int(0), int(1200), int(580)))
im.save("mapa3.png")

document = Document()
document.add_heading(avali, 2)
document.add_paragraph("PROGRESSO ATIVIDADE")
document.add_picture("mapa1.png", width=Inches(6.50))
document.add_paragraph("ASSUNTO EM DESTAQUE")
document.add_picture("mapa2.png", width=Inches(6.50))
document.add_paragraph("ALUNOS EM DESTAQUE")
document.add_picture("mapa3.png", width=Inches(6.50))
document.save("docteste.docx")

#driver.close()